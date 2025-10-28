from airflow import DAG
from airflow.decorators import task
from airflow.providers.http.hooks.http import HttpHook
from airflow.providers.postgres.hooks.postgres import PostgresHook
import requests
from datetime import datetime, timedelta
import os
from dotenv import load_dotenv
import pandas as pd
from io import StringIO
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Load environment variables
load_dotenv()
api_key = os.getenv('ALPHAVANTAGE_API_KEY')
POSTGRES_CONN_ID = 'postgres_default'

default_args = {
    'owner': 'airflow',
    'start_date': datetime(2025, 10, 20),
}

with DAG(
    dag_id='stock_prices_pipeline',
    default_args=default_args,
    schedule='@daily',
    catchup=False
) as dag:

    # ---------------------------
    # 1️⃣ Extract Task
    # ---------------------------
    @task()
    def extract_stock_prices_data():
        url = "https://www.alphavantage.co/query"
        params = {
            'function': 'TIME_SERIES_INTRADAY',
            'symbol': 'IBM',
            'interval': '60min',
            'outputsize': 'compact',
            'datatype': 'csv',
            'apikey': api_key
        }

        response = requests.get(url, params=params)
        if response.status_code == 200:
            df = pd.read_csv(StringIO(response.text))
            return df
        else:
            raise Exception("API request failed with status code:", response.status_code)

    # ---------------------------
    # 2️⃣ Load Task
    # ---------------------------
    @task()
    def load_data(df):
        pg_hook = PostgresHook(postgres_conn_id=POSTGRES_CONN_ID)
        conn = pg_hook.get_conn()
        cursor = conn.cursor()

        cursor.execute("""
        DROP TABLE IF EXISTS stock_prices CASCADE;
        CREATE TABLE IF NOT EXISTS stock_prices(
            time TIMESTAMP,
            open FLOAT,
            high FLOAT,
            low FLOAT,
            close FLOAT,
            volume INT
        );
        """)

        for _, row in df.iterrows():
            cursor.execute("""
                INSERT INTO stock_prices(time, open, high, low, close, volume)
                VALUES (%s, %s, %s, %s, %s, %s);
            """, (
                row['timestamp'],
                row['open'],
                row['high'],
                row['low'],
                row['close'],
                row['volume']
            ))

        conn.commit()
        cursor.close()
        conn.close()
        return True

    # ---------------------------
    # 3️⃣ Reporting Task
    # ---------------------------
    @task()
    def stock_prices_report():
        pg_hook = PostgresHook(postgres_conn_id=POSTGRES_CONN_ID)
        conn = pg_hook.get_conn()
        query = 'SELECT * FROM stock_prices'
        df = pd.read_sql(query, conn)
        conn.close()

        df['timestamp'] = pd.to_datetime(df['time'])
        df['date'] = df['timestamp'].dt.date

        base_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(base_dir, "reports")
        os.makedirs(output_dir, exist_ok=True)

        # ---------------------------
        # Plot 1: Close price over time
        # ---------------------------
        sns.set(style="whitegrid")
        plt.figure(figsize=(10, 5))
        sns.lineplot(x='timestamp', y='close', data=df)
        plt.title('Close Price Over Time')
        plt.xticks(rotation=45)
        plt.tight_layout()
        close_path = os.path.join(output_dir, "Close_price_over_time.png")
        plt.savefig(close_path, dpi=300)
        plt.close()

        # ---------------------------
        # Plot 2: Moving Average
        # ---------------------------
        df['ma_5'] = df['close'].rolling(5).mean()
        plt.figure(figsize=(10, 5))
        sns.lineplot(x='timestamp', y='close', data=df, label='Close')
        sns.lineplot(x='timestamp', y='ma_5', data=df, label='5-Period MA')
        plt.title('Close Price with Moving Average')
        plt.legend()
        plt.xticks(rotation=45)
        plt.tight_layout()
        ma_path = os.path.join(output_dir, "moving_average.png")
        plt.savefig(ma_path, dpi=300)
        plt.close()

        # ---------------------------
        # Summary Table
        # ---------------------------
        daily_summary = (
            df.groupby('date')
            .agg(
                open=('open', 'first'),
                high=('high', 'max'),
                low=('low', 'min'),
                close=('close', 'last'),
                volume=('volume', 'sum')
            )
            .reset_index()
        )
        last_day = daily_summary.iloc[-1]

        # ---------------------------
        # Create Word Report
        # ---------------------------
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = os.path.join(output_dir, f'Stock_Report_{timestamp}.docx')

        doc = Document()
        doc.add_heading("📊 Stock Market Report", level=1)
        doc.add_paragraph("This report contains summary information for the previous day.")

        doc.add_heading("Summary", level=2)
        doc.add_paragraph(f"Date: {last_day['date']}")
        doc.add_paragraph(f"Open: {last_day['open']:.2f}")
        doc.add_paragraph(f"Close: {last_day['close']:.2f}")
        doc.add_paragraph(f"Low: {last_day['low']:.2f}")
        doc.add_paragraph(f"High: {last_day['high']:.2f}")
        doc.add_paragraph(f"Volume: {int(last_day['volume'])}")

        doc.add_picture(ma_path, width=Inches(6))
        doc.add_picture(close_path, width=Inches(6))
        doc.save(report_path)

        print(f"✅ Report saved to: {report_path}")
        return report_path
        

    # ---------------------------
    # DAG Task Dependencies
    # ---------------------------
    stock_prices = extract_stock_prices_data()
    loaded = load_data(stock_prices)
    loaded >> stock_prices_report()
